#!/usr/bin/env python3
"""
Document Processor for Llamita Voice Assistant
Handles document upload, parsing, and storage for AI context
"""

import os
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Optional, Tuple
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import tkinter.ttk as ttk

# Document processing libraries
try:
    import PyPDF2
    # Test if we can create a PdfReader
    PDF_AVAILABLE = True
    print("✅ PyPDF2 imported successfully")
except ImportError as e:
    PDF_AVAILABLE = False
    print(f"❌ PyPDF2 import failed: {e}")

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

class DocumentProcessor:
    def __init__(self, storage_dir: str = "documents"):
        """
        Initialize the document processor
        
        Args:
            storage_dir: Directory to store processed documents
        """
        self.storage_dir = storage_dir
        self.documents = {}  # Store document metadata and content
        self.document_chunks = {}  # Store document chunks for context
        self.chunk_size = 1000  # Characters per chunk
        self.overlap = 200  # Overlap between chunks
        
        # Create storage directory if it doesn't exist
        os.makedirs(storage_dir, exist_ok=True)
        
        # Load existing documents (lazy loading to improve startup time)
        self._documents_loaded = False
    
    def load_documents(self):
        """Load existing documents from storage (optimized version)"""
        if self._documents_loaded:
            return
            
        metadata_file = os.path.join(self.storage_dir, "metadata.json")
        print(f"Loading documents from: {metadata_file}")
        
        if os.path.exists(metadata_file):
            try:
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    self.documents = json.load(f)
                
                print(f"Loaded {len(self.documents)} documents from metadata")
                
                # Load chunks asynchronously to avoid blocking the UI
                self._load_chunks_async()
                
            except Exception as e:
                print(f"Error loading documents: {e}")
        else:
            print("No metadata file found, starting with empty document list")
        
        self._documents_loaded = True
    
    def _load_chunks_async(self):
        """Load document chunks asynchronously to avoid blocking the UI"""
        def load_chunks_worker():
            try:
                for doc_id in self.documents:
                    chunks_file = os.path.join(self.storage_dir, f"{doc_id}_chunks.json")
                    if os.path.exists(chunks_file):
                        with open(chunks_file, 'r', encoding='utf-8') as f:
                            self.document_chunks[doc_id] = json.load(f)
                        print(f"Loaded chunks for document: {self.documents[doc_id].get('filename', 'Unknown')}")
                    else:
                        print(f"Warning: Chunks file not found for document: {doc_id}")
            except Exception as e:
                print(f"Error loading chunks: {e}")
        
        # Start loading chunks in background
        import threading
        thread = threading.Thread(target=load_chunks_worker, daemon=True)
        thread.start()
    
    def get_chunks_for_document(self, doc_id: str) -> List[Dict]:
        """Get chunks for a specific document, loading them if needed"""
        if doc_id not in self.document_chunks:
            # Load chunks for this specific document
            chunks_file = os.path.join(self.storage_dir, f"{doc_id}_chunks.json")
            if os.path.exists(chunks_file):
                try:
                    with open(chunks_file, 'r', encoding='utf-8') as f:
                        self.document_chunks[doc_id] = json.load(f)
                except Exception as e:
                    print(f"Error loading chunks for {doc_id}: {e}")
                    return []
            else:
                return []
        
        return self.document_chunks.get(doc_id, [])
    
    def _save_documents_async(self):
        """Save documents asynchronously to avoid blocking"""
        import threading
        
        def save_worker():
            try:
                # Save metadata
                metadata_file = os.path.join(self.storage_dir, "metadata.json")
                with open(metadata_file, 'w', encoding='utf-8') as f:
                    json.dump(self.documents, f, indent=2, ensure_ascii=False)
                
                # Save chunks for each document (create a copy to avoid iteration issues)
                doc_chunks_copy = dict(self.document_chunks)
                for doc_id, chunks in doc_chunks_copy.items():
                    chunks_file = os.path.join(self.storage_dir, f"{doc_id}_chunks.json")
                    with open(chunks_file, 'w', encoding='utf-8') as f:
                        json.dump(chunks, f, indent=2, ensure_ascii=False)
            except Exception as e:
                print(f"Error saving documents: {e}")
        
        # Start saving in background
        thread = threading.Thread(target=save_worker, daemon=True)
        thread.start()
    
    def save_documents(self):
        """Save document metadata and chunks to storage (synchronous version)"""
        try:
            # Save metadata
            metadata_file = os.path.join(self.storage_dir, "metadata.json")
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(self.documents, f, indent=2, ensure_ascii=False)
            
            # Save chunks for each document
            for doc_id, chunks in self.document_chunks.items():
                chunks_file = os.path.join(self.storage_dir, f"{doc_id}_chunks.json")
                with open(chunks_file, 'w', encoding='utf-8') as f:
                    json.dump(chunks, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving documents: {e}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats"""
        formats = [".txt"]
        
        if PDF_AVAILABLE:
            formats.append(".pdf")
        
        if DOCX_AVAILABLE:
            formats.append(".docx")
        
        if PANDAS_AVAILABLE:
            formats.extend([".csv", ".xlsx", ".xls"])
        
        return formats
    
    def process_document(self, file_path: str) -> Optional[str]:
        """
        Process a document and extract its text content (ultra-fast version)
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document ID if successful, None otherwise
        """
        try:
            if not os.path.exists(file_path):
                print(f"File not found: {file_path}")
                return None
            
            # Ultra-fast validation
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                print(f"File is empty: {file_path}")
                return None
            
            # Skip very large files immediately
            if file_size > 10 * 1024 * 1024:  # 10MB limit
                print(f"File too large: {file_size / 1024 / 1024:.1f}MB")
                return None
            
            # Generate document ID
            doc_id = self._generate_doc_id(file_path)
            
            # Extract text with shorter timeout for faster processing
            text_content = self._extract_text_with_timeout(file_path, timeout=15)
            if not text_content:
                print(f"Could not extract text from: {file_path}")
                return None
            
            # Create document metadata
            doc_metadata = {
                "filename": os.path.basename(file_path),
                "filepath": file_path,
                "size": file_size,
                "uploaded_at": datetime.now().isoformat(),
                "content_length": len(text_content),
                "chunks_count": 0
            }
            
            # Store document
            self.documents[doc_id] = doc_metadata
            
            # Create chunks with ultra-fast processing
            chunks = self._create_chunks_ultra_fast(text_content)
            self.document_chunks[doc_id] = chunks
            doc_metadata["chunks_count"] = len(chunks)
            
            # Save to storage (synchronous to ensure it's saved)
            self.save_documents()
            
            print(f"Successfully processed document: {os.path.basename(file_path)}")
            print(f"Created {len(chunks)} chunks for context")
            print(f"Document saved with ID: {doc_id}")
            
            return doc_id
            
        except Exception as e:
            print(f"Error processing document {file_path}: {e}")
            return None
    
    def _generate_doc_id(self, file_path: str) -> str:
        """Generate a unique document ID based on file content"""
        try:
            with open(file_path, 'rb') as f:
                content = f.read()
                return hashlib.md5(content).hexdigest()[:16]
        except:
            # Fallback to filename-based ID
            return hashlib.md5(file_path.encode()).hexdigest()[:16]
    
    def _extract_text_with_timeout(self, file_path: str, timeout: int = 15) -> Optional[str]:
        """Extract text content with timeout to prevent hanging (ultra-fast version)"""
        import threading
        import queue
        
        result_queue = queue.Queue()
        
        def extract_worker():
            try:
                file_ext = os.path.splitext(file_path)[1].lower()
                
                # Ultra-fast text extraction
                if file_ext == ".txt":
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        # Limit content size for speed
                        if len(content) > 100000:  # 100KB limit
                            content = content[:100000]
                        result_queue.put(content)
                
                elif file_ext == ".pdf" and PDF_AVAILABLE:
                    print(f"Processing PDF: {file_path}")
                    content = self._extract_pdf_text_fast(file_path)
                    if not content.strip():
                        print(f"No text extracted from PDF: {file_path}")
                        result_queue.put(None)
                    else:
                        print(f"Successfully extracted {len(content)} characters from PDF")
                        result_queue.put(content)
                elif file_ext == ".pdf" and not PDF_AVAILABLE:
                    print(f"PDF processing not available - PDF_AVAILABLE: {PDF_AVAILABLE}")
                    result_queue.put(None)
                
                elif file_ext == ".docx" and DOCX_AVAILABLE:
                    content = self._extract_docx_text_fast(file_path)
                    result_queue.put(content)
                
                elif file_ext in [".csv", ".xlsx", ".xls"] and PANDAS_AVAILABLE:
                    content = self._extract_spreadsheet_text_fast(file_path)
                    result_queue.put(content)
                
                else:
                    print(f"Unsupported file format: {file_ext}")
                    result_queue.put(None)
                    
            except Exception as e:
                print(f"Error extracting text from {file_path}: {e}")
                result_queue.put(None)
        
        # Start extraction in thread
        thread = threading.Thread(target=extract_worker, daemon=True)
        thread.start()
        
        # Wait for result with shorter timeout
        try:
            result = result_queue.get(timeout=timeout)
            return result
        except queue.Empty:
            print(f"Text extraction timed out for {file_path}")
            return None
    
    def _extract_text(self, file_path: str) -> Optional[str]:
        """Extract text content from various document formats (legacy method)"""
        return self._extract_text_with_timeout(file_path)
    
    def _extract_pdf_text_fast(self, file_path: str) -> str:
        """Extract text from PDF file (ultra-fast version)"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if hasattr(pdf_reader, 'is_encrypted') and pdf_reader.is_encrypted:
                    print(f"PDF is encrypted: {file_path}")
                    return text
                
                # Limit to first 10 pages for speed
                max_pages = min(10, len(pdf_reader.pages))
                print(f"Processing {max_pages} pages from PDF")
                
                for i in range(max_pages):
                    try:
                        page = pdf_reader.pages[i]
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            print(f"Extracted {len(page_text)} characters from page {i+1}")
                    except Exception as e:
                        print(f"Error extracting text from page {i}: {e}")
                        continue
                    
                    # Limit total text size
                    if len(text) > 50000:  # 50KB limit
                        text = text[:50000]
                        break
                        
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
        return text
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file (legacy method)"""
        return self._extract_pdf_text_fast(file_path)
    
    def _extract_docx_text_fast(self, file_path: str) -> str:
        """Extract text from DOCX file (ultra-fast version)"""
        text = ""
        try:
            doc = docx.Document(file_path)
            # Limit paragraphs for speed
            max_paragraphs = min(100, len(doc.paragraphs))
            for i in range(max_paragraphs):
                text += doc.paragraphs[i].text + "\n"
                # Limit total text size
                if len(text) > 50000:  # 50KB limit
                    text = text[:50000]
                    break
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file (legacy method)"""
        return self._extract_docx_text_fast(file_path)
    
    def _extract_spreadsheet_text_fast(self, file_path: str) -> str:
        """Extract text from spreadsheet files (ultra-fast version)"""
        text = ""
        try:
            # Read only first 100 rows for speed
            if file_path.endswith(('.xlsx', '.xls')):
                df = pd.read_excel(file_path, nrows=100)
            else:
                df = pd.read_csv(file_path, nrows=100)
            
            # Convert to string with size limit
            text = df.to_string(index=False)
            if len(text) > 50000:  # 50KB limit
                text = text[:50000]
        except Exception as e:
            print(f"Error extracting spreadsheet text: {e}")
        return text
    
    def _extract_spreadsheet_text(self, file_path: str) -> str:
        """Extract text from spreadsheet files (legacy method)"""
        return self._extract_spreadsheet_text_fast(file_path)
    
    def _create_chunks_ultra_fast(self, text: str) -> List[Dict]:
        """Create overlapping chunks from text with ultra-fast performance"""
        chunks = []
        start = 0
        text_length = len(text)
        
        # Ultra-fast chunking with smaller chunks and fewer limits
        max_chunks = 50  # Reduced for speed
        chunk_size = 800  # Smaller chunks for faster processing
        overlap = 100     # Reduced overlap
        
        while start < text_length and len(chunks) < max_chunks:
            end = start + chunk_size
            
            # Simple boundary detection for speed
            if end < text_length:
                # Quick sentence boundary search
                for i in range(min(end, text_length - 1), max(start, end - 30), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "start": start,
                    "end": end,
                    "length": len(chunk_text)
                })
            
            # Move start position with reduced overlap
            start = end - overlap
            if start >= text_length:
                break
        
        return chunks
    
    def _create_chunks_optimized(self, text: str) -> List[Dict]:
        """Create overlapping chunks from text with optimized performance"""
        chunks = []
        start = 0
        text_length = len(text)
        
        # Limit chunks for very large documents
        max_chunks = 100
        
        while start < text_length and len(chunks) < max_chunks:
            end = start + self.chunk_size
            
            # Try to break at a sentence boundary
            if end < text_length:
                # Look for sentence endings (optimized search)
                for i in range(min(end, text_length - 1), max(start, end - 50), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append({
                    "text": chunk_text,
                    "start": start,
                    "end": end,
                    "length": len(chunk_text)
                })
            
            # Move start position with overlap
            start = end - self.overlap
            if start >= text_length:
                break
        
        return chunks
    
    def _create_chunks(self, text: str) -> List[Dict]:
        """Create overlapping chunks from text for better context (legacy method)"""
        return self._create_chunks_optimized(text)
    
    def get_document_context(self, query: str, max_chunks: int = 3) -> str:
        """
        Get relevant document chunks for a query
        
        Args:
            query: The user's query
            max_chunks: Maximum number of chunks to return
            
        Returns:
            Relevant document context as string
        """
        if not self.documents:
            return ""
        
        # Simple keyword matching for now
        # In a more advanced implementation, you could use embeddings or semantic search
        query_words = query.lower().split()
        relevant_chunks = []
        
        for doc_id in self.documents:
            doc_info = self.documents.get(doc_id, {})
            filename = doc_info.get("filename", "Unknown")
            
            # Load chunks for this document on demand
            chunks = self.get_chunks_for_document(doc_id)
            
            for chunk in chunks:
                chunk_text = chunk["text"].lower()
                score = 0
                
                # Simple keyword scoring
                for word in query_words:
                    if word in chunk_text:
                        score += 1
                
                if score > 0:
                    relevant_chunks.append({
                        "doc_id": doc_id,
                        "filename": filename,
                        "chunk": chunk,
                        "score": score
                    })
        
        # Sort by relevance score and take top chunks
        relevant_chunks.sort(key=lambda x: x["score"], reverse=True)
        
        # Build context string
        context_parts = []
        for item in relevant_chunks[:max_chunks]:
            context_parts.append(f"From document '{item['filename']}':\n{item['chunk']['text']}\n")
        
        return "\n".join(context_parts)
    
    def list_documents(self) -> List[Dict]:
        """Get list of all uploaded documents"""
        # Ensure documents are loaded
        self.load_documents()
        
        documents_list = [
            {
                "id": doc_id,
                "filename": metadata.get("filename", "Unknown"),
                "uploaded_at": metadata.get("uploaded_at", ""),
                "size": metadata.get("size", 0),
                "chunks_count": metadata.get("chunks_count", 0)
            }
            for doc_id, metadata in self.documents.items()
        ]
        
        print(f"Returning {len(documents_list)} documents from list_documents")
        for doc in documents_list:
            print(f"  - {doc['filename']} ({doc['chunks_count']} chunks)")
        
        return documents_list
    
    def remove_document(self, doc_id: str) -> bool:
        """Remove a document from storage"""
        try:
            if doc_id in self.documents:
                del self.documents[doc_id]
            
            if doc_id in self.document_chunks:
                del self.document_chunks[doc_id]
            
            # Remove chunk file
            chunks_file = os.path.join(self.storage_dir, f"{doc_id}_chunks.json")
            if os.path.exists(chunks_file):
                os.remove(chunks_file)
            
            self.save_documents()
            return True
        except Exception as e:
            print(f"Error removing document {doc_id}: {e}")
            return False
    
    def remove_multiple_documents(self, doc_ids: List[str]) -> Dict[str, bool]:
        """Remove multiple documents and return success status for each"""
        results = {}
        for doc_id in doc_ids:
            results[doc_id] = self.remove_document(doc_id)
        return results
    
    def get_document_info(self, doc_id: str) -> Optional[Dict]:
        """Get detailed information about a document"""
        if doc_id not in self.documents:
            return None
        
        doc_info = self.documents[doc_id].copy()
        doc_info['chunks_count'] = len(self.document_chunks.get(doc_id, []))
        doc_info['storage_size'] = self._get_document_storage_size(doc_id)
        return doc_info
    
    def _get_document_storage_size(self, doc_id: str) -> int:
        """Get the storage size of a document in bytes"""
        try:
            chunks_file = os.path.join(self.storage_dir, f"{doc_id}_chunks.json")
            if os.path.exists(chunks_file):
                return os.path.getsize(chunks_file)
            return 0
        except:
            return 0
    
    def get_storage_stats(self) -> Dict:
        """Get storage statistics"""
        total_docs = len(self.documents)
        
        # Calculate total chunks and size more efficiently
        total_chunks = 0
        total_size = 0
        
        for doc_id in self.documents:
            chunks = self.get_chunks_for_document(doc_id)
            total_chunks += len(chunks)
            total_size += self._get_document_storage_size(doc_id)
        
        return {
            'total_documents': total_docs,
            'total_chunks': total_chunks,
            'total_size_bytes': total_size,
            'total_size_mb': round(total_size / (1024 * 1024), 2)
        }
    
    def clear_all_documents(self):
        """Remove all documents"""
        self.documents.clear()
        self.document_chunks.clear()
        self.save_documents()
        
        # Clean up chunk files
        try:
            for filename in os.listdir(self.storage_dir):
                if filename.endswith('_chunks.json'):
                    os.remove(os.path.join(self.storage_dir, filename))
        except Exception as e:
            print(f"Error cleaning up chunk files: {e}")
    
    def cleanup_orphaned_files(self):
        """Remove chunk files for documents that no longer exist in metadata"""
        try:
            for filename in os.listdir(self.storage_dir):
                if filename.endswith('_chunks.json'):
                    doc_id = filename.replace('_chunks.json', '')
                    if doc_id not in self.documents:
                        os.remove(os.path.join(self.storage_dir, filename))
                        print(f"Removed orphaned chunk file: {filename}")
        except Exception as e:
            print(f"Error cleaning up orphaned files: {e}")


class DocumentUploadDialog:
    """GUI dialog for document upload"""
    
    def __init__(self, parent, document_processor: DocumentProcessor):
        self.parent = parent
        self.document_processor = document_processor
        self.result = None
        
        # Create dialog window with optimized settings
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Documents")
        self.dialog.geometry("750x900")  # Much larger size to show all content without scrolling
        self.dialog.transient(parent)
        
        # Ensure dialog is visible and properly configured
        self.dialog.deiconify()  # Make sure dialog is visible
        self.dialog.lift()  # Bring to front
        self.dialog.focus_force()  # Force focus
        
        # Make dialog modal to ensure it's properly displayed
        self.dialog.grab_set()
        
        # Center the dialog
        self.dialog.geometry("+%d+%d" % (
            parent.winfo_rootx() + parent.winfo_width()//2 - 375,
            parent.winfo_rooty() + parent.winfo_height()//2 - 450
        ))
        
        # Setup UI directly without scrollable frame
        self.setup_ui()
        
        # Focus on dialog
        self.dialog.focus_set()
        
        # Bind escape key to close
        self.dialog.bind('<Escape>', lambda e: self.dialog.destroy())
        
        # Ensure dialog is visible
        self.dialog.deiconify()
        self.dialog.lift()
        
        # Wait for dialog to close
        self.dialog.wait_window()
    
    def create_scrollable_frame(self):
        """Create a scrollable frame for the dialog content"""
        # Create main container
        self.main_container = tk.Frame(self.dialog)
        self.main_container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Create canvas with scrollbar
        self.canvas = tk.Canvas(self.main_container, bg='white')
        self.scrollbar = ttk.Scrollbar(self.main_container, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = tk.Frame(self.canvas)
        
        # Configure canvas
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        
        # Pack canvas and scrollbar
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")
        
        # Bind mouse wheel to scroll
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)
    
    def _on_mousewheel(self, event):
        """Handle mouse wheel scrolling"""
        self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")
    
    def setup_ui(self):
        """Setup the upload dialog UI"""
        # Main frame
        main_frame = tk.Frame(self.dialog)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        # Title
        title_label = tk.Label(
            main_frame,
            text="Documents",
            font=("Helvetica", 16, "bold")
        )
        title_label.pack(pady=(0, 20))
        
        # Storage stats (load immediately for faster response)
        try:
            stats = self.document_processor.get_storage_stats()
            stats_text = f"📊 Storage: {stats['total_documents']} docs, {stats['total_chunks']} chunks, {stats['total_size_mb']} MB"
        except Exception as e:
            stats_text = "📊 Storage: Loading..."
        
        self.stats_label = tk.Label(
            main_frame,
            text=stats_text,
            font=("Helvetica", 10),
            fg="gray"
        )
        self.stats_label.pack(pady=(0, 20))
        
        # Load documents immediately for faster response
        self.load_documents()
        
        # Local Files Section
        local_frame = tk.LabelFrame(main_frame, text="💻 Local Files", font=("Helvetica", 12, "bold"), padx=15, pady=15)
        local_frame.pack(fill=tk.X, pady=(0, 15))
        
        # Supported formats
        formats_label = tk.Label(local_frame, text="Supported formats: .txt, .pdf, .docx", font=("Helvetica", 9), fg="gray")
        formats_label.pack(anchor=tk.W, pady=(0, 10))
        
        # File selection
        file_frame = tk.Frame(local_frame)
        file_frame.pack(fill=tk.X, pady=(0, 10))
        
        self.file_entry = tk.Entry(file_frame, font=("Helvetica", 10))
        self.file_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        
        browse_button = tk.Button(
            file_frame,
            text="Browse",
            command=self.browse_file,
            font=("Helvetica", 10),
            bg="#FF9800",
            fg="white",
            relief=tk.RAISED,
            padx=15,
            pady=5
        )
        browse_button.pack(side=tk.RIGHT)
        
        # Upload button
        self.upload_button = tk.Button(
            local_frame,
            text="Upload Document",
            command=self.upload_document,
            font=("Helvetica", 10, "bold"),
            bg="#4CAF50",
            fg="white",
            relief=tk.RAISED,
            padx=20,
            pady=8
        )
        self.upload_button.pack(pady=(5, 0))
        
        # Status label for upload feedback
        self.status_label = tk.Label(
            local_frame,
            text="Select a document to upload",
            font=("Helvetica", 9),
            fg="gray"
        )
        self.status_label.pack(anchor=tk.W, pady=(5, 0))
        
        # Uploaded Documents Section
        docs_frame = tk.LabelFrame(main_frame, text="📚 Uploaded Documents", font=("Helvetica", 12, "bold"), padx=15, pady=15)
        docs_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        
        # Documents list
        self.documents_listbox = tk.Listbox(
            docs_frame,
            font=("Helvetica", 10),
            selectmode=tk.SINGLE,
            height=12  # Increased height for better visibility
        )
        self.documents_listbox.pack(fill=tk.BOTH, expand=True, pady=(0, 10))
        
        # Document info label
        self.doc_info_label = tk.Label(docs_frame, text="Select a document to view details", font=("Helvetica", 9), fg="gray")
        self.doc_info_label.pack(anchor=tk.W, pady=(0, 10))
        
        # Action buttons frame
        action_frame = tk.Frame(docs_frame)
        action_frame.pack(fill=tk.X)
        
        # Remove selected button
        remove_button = tk.Button(
            action_frame,
            text="🗑️ Remove Selected",
            command=self.remove_selected_document,
            font=("Helvetica", 9),
            bg="#F44336",
            fg="white",
            relief=tk.RAISED,
            padx=10,
            pady=5
        )
        remove_button.pack(side=tk.LEFT, padx=(0, 10))
        
        # Remove all button
        remove_all_button = tk.Button(
            action_frame,
            text="🗑️ Remove All",
            command=self.remove_all_documents,
            font=("Helvetica", 9),
            bg="#FF5722",
            fg="white",
            relief=tk.RAISED,
            padx=10,
            pady=5
        )
        remove_all_button.pack(side=tk.LEFT, padx=(0, 10))
        
        # Cleanup button
        cleanup_button = tk.Button(
            action_frame,
            text="🧹 Cleanup",
            command=self.cleanup_orphaned_files,
            font=("Helvetica", 9),
            bg="#9C27B0",
            fg="white",
            relief=tk.RAISED,
            padx=10,
            pady=5
        )
        cleanup_button.pack(side=tk.LEFT)
        
        # Close button
        close_button = tk.Button(
            main_frame,
            text="Close",
            command=self.dialog.destroy,
            font=("Helvetica", 10, "bold"),
            bg="#607D8B",
            fg="white",
            relief=tk.RAISED,
            padx=20,
            pady=8
        )
        close_button.pack(pady=(10, 0))
        
        # Bind double-click for document info
        self.documents_listbox.bind('<Double-Button-1>', self.show_document_info)
    
    def load_documents(self):
        """Load and display documents in the listbox (optimized)"""
        try:
            if self.document_processor:
                # Update document list immediately (fast)
                self.update_document_list()
                
                # Update stats asynchronously to avoid blocking
                self.dialog.after(100, self.update_storage_stats)
        except Exception as e:
            print(f"Error loading documents: {e}")
            # Update UI to show error state
            try:
                self.stats_label.config(text="📊 Storage: Error loading stats")
            except:
                pass
    
    def update_storage_stats(self):
        """Update storage statistics display (optimized)"""
        try:
            # Get basic stats first (fast)
            total_docs = len(self.document_processor.documents)
            stats_text = f"📊 Storage: {total_docs} docs, loading..."
            self.stats_label.config(text=stats_text)
            
            # Update with full stats asynchronously
            def update_full_stats():
                try:
                    stats = self.document_processor.get_storage_stats()
                    stats_text = f"📊 Storage: {stats['total_documents']} docs, {stats['total_chunks']} chunks, {stats['total_size_mb']} MB"
                    self.stats_label.config(text=stats_text)
                except Exception as e:
                    self.stats_label.config(text="📊 Storage: Unable to load stats")
            
            # Run full stats update in background
            import threading
            thread = threading.Thread(target=update_full_stats, daemon=True)
            thread.start()
            
        except Exception as e:
            self.stats_label.config(text="📊 Storage: Unable to load stats")
    
    def show_document_info(self, event=None):
        """Show detailed information about the selected document"""
        selection = self.documents_listbox.curselection()
        if not selection:
            return
        
        documents = self.document_processor.list_documents()
        if selection[0] < len(documents):
            doc_id = documents[selection[0]]["id"]
            doc_info = self.document_processor.get_document_info(doc_id)
            
            if doc_info:
                info_text = f"📄 {doc_info['filename']}\n"
                info_text += f"📅 Uploaded: {doc_info['uploaded_at']}\n" # Changed from upload_date to uploaded_at
                info_text += f"📊 Chunks: {doc_info['chunks_count']}\n"
                info_text += f"💾 Size: {round(doc_info['storage_size'] / 1024, 1)} KB\n"
                info_text += f"🔍 Type: {doc_info.get('file_type', 'Unknown')}"
                
                self.doc_info_label.config(text=info_text)
    
    def remove_all_documents(self):
        """Remove all documents with confirmation"""
        if messagebox.askyesno("Remove All Documents", 
                             "Are you sure you want to remove ALL documents?\n\nThis action cannot be undone."):
            self.document_processor.clear_all_documents()
            self.update_document_list()
            self.update_storage_stats()
            self.doc_info_label.config(text="All documents removed")
            # The status_label was removed from setup_ui, so this line is no longer applicable.
            # self.status_label.config(text="✅ All documents removed successfully")
    
    def cleanup_orphaned_files(self):
        """Clean up orphaned chunk files"""
        try:
            self.document_processor.cleanup_orphaned_files()
            # The status_label was removed from setup_ui, so this line is no longer applicable.
            # self.status_label.config(text="✅ Cleanup completed")
        except Exception as e:
            # The status_label was removed from setup_ui, so this line is no longer applicable.
            # self.status_label.config(text=f"❌ Cleanup failed: {str(e)[:50]}")
            pass # No status_label to update
    
    def update_supported_formats(self):
        """Update the supported formats display"""
        try:
            formats = self.document_processor.get_supported_formats()
            formats_text = "Supported formats: " + ", ".join(formats)
            # The formats_label was removed from setup_ui, so this line is no longer applicable.
            # self.formats_label.config(text=formats_text)
        except Exception as e:
            # The formats_label was removed from setup_ui, so this line is no longer applicable.
            # self.formats_label.config(text="Supported formats: .txt, .pdf, .docx, .csv, .xlsx, .xls")
            pass # No formats_label to update
    
    def browse_file(self):
        """Open file browser dialog"""
        # Use a more efficient approach for macOS
        try:
            # Show a simple message first
            # The status_label was removed from setup_ui, so this line is no longer applicable.
            # self.status_label.config(text="Opening file browser...")
            self.dialog.update()
            
            # Try to use a more modern file dialog if available
            try:
                file_path = filedialog.askopenfilename(
                    title="Select Document",
                    filetypes=[
                        ("All supported", "*.txt;*.pdf;*.docx;*.csv;*.xlsx;*.xls"),
                        ("Text files", "*.txt"),
                        ("PDF files", "*.pdf"),
                        ("Word documents", "*.docx"),
                        ("Excel files", "*.xlsx;*.xls"),
                        ("CSV files", "*.csv"),
                        ("All files", "*.*")
                    ]
                )
            except Exception as e:
                print(f"File dialog error: {e}")
                # Fallback to basic file dialog
                file_path = filedialog.askopenfilename(title="Select Document")
            
            if file_path:
                # Validate file exists
                if os.path.exists(file_path):
                    self.file_entry.delete(0, tk.END)
                    self.file_entry.insert(0, file_path)
                    self.upload_button.config(state="normal")
                    self.status_label.config(text=f"Selected: {os.path.basename(file_path)}")
        except Exception as e:
            print(f"Browse file error: {e}")
            messagebox.showerror("Error", f"Error opening file browser: {str(e)}")
    
    def upload_document(self):
        """Upload the selected document"""
        file_path = self.file_entry.get()
        if not file_path:
            return
        
        # Validate file exists before processing
        if not os.path.exists(file_path):
            self.status_label.config(text="❌ File not found")
            return
        
        self.status_label.config(text="Processing document...")
        self.upload_button.config(state="disabled")
        
        # Process document in a separate thread with ultra-fast processing
        def process_thread():
            try:
                # Ultra-fast file size check
                file_size = os.path.getsize(file_path)
                if file_size > 10 * 1024 * 1024:  # 10MB limit for speed
                    self.parent.after(0, lambda: messagebox.showwarning(
                        "File Too Large",
                        "File too large (max 10MB for speed). Please choose a smaller file."
                    ))
                    return
                
                # Update status during processing
                self.parent.after(0, lambda: messagebox.showinfo(
                    "Uploading Document",
                    f"Uploading {os.path.basename(file_path)}..."
                ))
                
                # Process with timeout
                import threading
                import queue
                
                result_queue = queue.Queue()
                
                def process_worker():
                    try:
                        doc_id = self.document_processor.process_document(file_path)
                        result_queue.put(doc_id)
                    except Exception as e:
                        result_queue.put(None)
                
                # Start processing with timeout
                worker_thread = threading.Thread(target=process_worker, daemon=True)
                worker_thread.start()
                
                # Wait for result with timeout
                try:
                    doc_id = result_queue.get(timeout=20)  # 20 second timeout
                    if doc_id:
                        self.parent.after(0, lambda: messagebox.showinfo(
                            "Upload Successful",
                            f"✅ Successfully uploaded: {os.path.basename(file_path)}"
                        ))
                        self.update_document_list()
                    else:
                        self.parent.after(0, lambda: messagebox.showerror(
                            "Upload Failed",
                            "Failed to process document"
                        ))
                except queue.Empty:
                    self.parent.after(0, lambda: messagebox.showerror(
                        "Processing Timed Out",
                        "Document processing timed out. Please try again."
                    ))
                    
            except Exception as e:
                self.parent.after(0, lambda: messagebox.showerror(
                    "Error",
                    f"Error: {str(e)[:50]}..."
                ))
            finally:
                self.parent.after(0, lambda: self.upload_button.config(state="normal"))
        
        # Use ultra-fast threading approach
        import threading
        thread = threading.Thread(target=process_thread, daemon=True)
        thread.start()
    
    def update_document_list(self):
        """Update the document list display"""
        self.documents_listbox.delete(0, tk.END)
        documents = self.document_processor.list_documents()
        
        for doc in documents:
            display_text = f"{doc['filename']} ({doc['chunks_count']} chunks)"
            self.documents_listbox.insert(tk.END, display_text)
    
    def remove_selected_document(self):
        """Remove the selected document"""
        selection = self.documents_listbox.curselection()
        if not selection:
            return
        
        documents = self.document_processor.list_documents()
        if selection[0] < len(documents):
            doc_id = documents[selection[0]]["id"]
            
            if messagebox.askyesno("Remove Document", 
                                 f"Are you sure you want to remove '{documents[selection[0]]['filename']}'?"):
                if self.document_processor.remove_document(doc_id):
                    self.update_document_list()
                    # The status_label was removed from setup_ui, so this line is no longer applicable.
                    # self.status_label.config(text="Document removed successfully")
    
    def close_dialog(self):
        """Close the upload dialog"""
        self.dialog.destroy()
